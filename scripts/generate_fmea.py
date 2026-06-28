#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PFMEA/DFMEA 报告生成器
======================

基于 AIAG & VDA FMEA 手册（2019 版）七步法，生成 FMEA 分析报告（.xlsx + .docx）。

依赖：
    - openpyxl（生成 Excel）
    - python-docx（生成 Word）

未安装时自动 pip install。

用法：
    python3 generate_fmea.py \\
        --fmea-type PFMEA \\
        --product "前保险杠总成（注塑件）" \\
        --customer "XX 汽车有限公司" \\
        --process-name "注塑成型" \\
        --process-steps "原料干燥,注塑成型,去毛刺,外观检验,包装" \\
        --template generic-fmea \\
        --auto-fill \\
        --output-dir ~/Desktop

⚠️ 评分基准（详见 references/sod_scoring_tables.md）：
    - S（严重度）：1-10，仅 S=10 用于"影响行车安全"
    - O（频度）：1-10，基于预防控制有效性
    - D（探测度）：1-10，基于探测方法成熟度
    - AP（行动优先级）：H/M/L，基于 S/O/D 组合查表
    - 严禁使用 RPN（=S×O×D），2019 版已废弃
"""

# ===== Windows 编码修复（必须在最前面） =====
# Windows 服务器默认 stdout 编码是 GBK/CP936，会导致中文输出乱码
# 强制设置为 UTF-8，确保 subprocess 调用时能正确解析中文输出
import sys
import io
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
if hasattr(sys.stderr, 'buffer'):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace', line_buffering=True)

import argparse
import json
import os
import subprocess
import datetime
from pathlib import Path

# ============================================================
# 依赖检查与自动安装
# ============================================================

REQUIRED_PACKAGES = {
    "openpyxl": "openpyxl",
    "docx": "python-docx",  # 注意：导入名是 docx，包名是 python-docx
}


def ensure_packages():
    """检查并自动安装所需的 Python 包。"""
    import importlib

    for import_name, pip_name in REQUIRED_PACKAGES.items():
        try:
            importlib.import_module(import_name)
        except ImportError:
            print(f"[INFO] {pip_name} 未安装，正在自动安装...")
            try:
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", pip_name, "--quiet"]
                )
            except subprocess.CalledProcessError as e:
                print(f"[ERROR] 安装 {pip_name} 失败：{e}")
                print(f"请手动执行：pip install {pip_name}")
                sys.exit(1)


ensure_packages()

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ============================================================
# 常量（FMEA 专用配色，无 8D 配色污染）
# ============================================================

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent  # pfmea-dfmea-skill/
TEMPLATES_DIR = SKILL_ROOT / "templates"

# FMEA 报告主色（深蓝系列，区别于 8D 的蓝黄色）
HEADER_FILL = "1F4E79"          # 深蓝色表头（FMEA 主色）
SUBHEADER_FILL = "2E75B6"       # 次级表头（亮蓝）
ALT_ROW_FILL = "E7EEF7"         # 交替行浅蓝
SECTION_TITLE_FILL = "FFF8E1"   # 章节标题淡黄色底

# AP 行动优先级专用配色（H=红/M=黄/L=绿，符合 2019 版手册惯例）
AP_H_FILL = "FF6B6B"            # 高优先级 = 红
AP_M_FILL = "FFD93D"            # 中优先级 = 黄
AP_L_FILL = "6BCB77"            # 低优先级 = 绿

# 特殊特性专用高亮（CC/SC）
CC_FILL = "FFC7CE"              # 关键特性（红色淡底）
SC_FILL = "FFEB9C"              # 特殊特性（黄色淡底）

# 措施状态专用配色
STATUS_OPEN_FILL = "FFE699"     # 已建议/进行中（黄）
STATUS_DONE_FILL = "C6EFCE"     # 已实施/已关闭（绿）

# 字体
FONT_NAME = "微软雅黑"
FONT_NAME_EN = "Microsoft YaHei"

# FMEA 团队角色化名（区别于 8D 的角色定义）
# FMEA 团队由手册第 1.5.3 节定义：推进者、设计工程师、过程工程师、质量工程师、测试工程师等
ROLE_NAMES = {
    "FMEA 推进者": "张伟",
    "设计工程师": "李娜",
    "过程工程师": "刘强",
    "质量工程师": "陈静",
    "测试工程师": "赵磊",
    "系统工程师": "周敏",
    "质量经理（审核）": "王芳",
    "项目经理": "孙健",
    "编制（FMEA 推进者）": "张伟",
    "审核（质量经理）": "王芳",
    "批准（项目经理）": "孙健",
}

# 部门对应的化名（用于优化措施责任人列）
DEPT_PERSON = {
    "设计部": "李娜",
    "工艺部": "刘强",
    "质量部": "陈静",
    "测试部": "赵磊",
    "系统部": "周敏",
    "项目部": "孙健",
    "生产部": "吴洋",
}

# 措施责任人轮换池（按 AP=H 优先 → M → L 分配）
PERSON_ROTATION = ["张伟", "李娜", "刘强", "陈静", "赵磊", "周敏", "王芳"]

# 内部分机号映射（仅用于 FMEA 团队表）
EXT_MAP = {
    "张伟": "8001", "李娜": "8002", "刘强": "8003", "陈静": "8004",
    "赵磊": "8005", "周敏": "8006", "王芳": "8007", "孙健": "8008", "吴洋": "8009",
}


# ============================================================
# FMEA 评分逻辑（2019 版 AP 矩阵，1000 种组合）
# ============================================================

def get_ap_priority(s: int, o: int, d: int) -> str:
    """根据 S/O/D 评级返回 AP 行动优先级（H/M/L）。

    基于 AIAG & VDA FMEA 手册（2019 版）AP 矩阵的简化实现。
    完整 1000 种组合详见 references/sod_scoring_tables.md。
    """
    # 边界保护
    s = max(1, min(10, int(s)))
    o = max(1, min(10, int(o)))
    d = max(1, min(10, int(d)))

    # S=1：AP 必为 L
    if s == 1:
        return "L"

    # S=9-10：通常为 H
    if s >= 9:
        if o == 1 and d == 1:
            return "M"
        return "H"

    # S=6-8：中等严重度
    if s >= 6:
        if o >= 6:
            return "H"
        if o >= 4 and d >= 7:
            return "H"
        if o <= 3 and d <= 3:
            return "L"
        return "M"

    # S=4-5：低严重度
    if s >= 4:
        if o >= 6 and d >= 7:
            return "H"
        if o >= 4 and d >= 7:
            return "M"
        return "L"

    # S=2-3：很低严重度
    if s >= 2:
        if o >= 4 and d >= 5:
            return "M"
        return "L"

    return "L"


def get_special_characteristic(s: int, ap: str) -> str:
    """根据 S 评分和 AP 判定特殊特性类型（CC/SC/空）。

    CC（Critical Characteristic）：S≥9 的关键特性
    SC（Significant Characteristic）：S=8 且 AP=H/M 的重要特性
    """
    if s >= 9:
        return "CC"
    if s == 8 and ap in ("H", "M"):
        return "SC"
    return ""


def check_sod_consistency(chain: dict) -> list:
    """检查 S/O/D 评分的逻辑一致性，返回警告列表。

    遵循 SKILL.md 第十章「行业常识基准」：
    - 严禁 O=10 配合"已使用 10 年的成熟产品"
    - 严禁 D=10 配合"100% 在线自动检测"
    - 严禁 S=10/O=10/D=10/AP=L 这种荒谬组合
    """
    warnings = []
    s = chain.get("s", 0)
    o = chain.get("o", 0)
    d = chain.get("d", 0)
    ap = chain.get("ap", "")

    # S=9-10 时 AP 必为 H（除非 O=1 且 D=1）
    if s >= 9 and ap != "H":
        if not (o == 1 and d == 1):
            warnings.append(f"S={s} 但 AP={ap}，应为 H（除非 O=1 且 D=1）")

    # S=1 时 AP 必为 L
    if s == 1 and ap != "L":
        warnings.append(f"S=1 但 AP={ap}，应为 L")

    # AP 与 S/O/D 组合一致性
    expected_ap = get_ap_priority(s, o, d)
    if ap and ap != expected_ap:
        warnings.append(f"AP={ap} 与 S/O/D={s}/{o}/{d} 不一致，应为 {expected_ap}")

    return warnings


# ============================================================
# 模板加载与占位符替换
# ============================================================

def load_template(template_slug: str, templates_dir: Path) -> dict:
    """加载模板 JSON。"""
    template_path = templates_dir / template_slug / "template.json"
    if not template_path.exists():
        print(f"[WARN] 模板不存在：{template_path}，回退到 generic-fmea")
        template_path = templates_dir / "generic-fmea" / "template.json"
        if not template_path.exists():
            raise FileNotFoundError(f"模板文件不存在：{template_path}")
    with open(template_path, "r", encoding="utf-8") as f:
        return json.load(f)


def substitute_placeholders(text: str, context: dict) -> str:
    """替换文本中的占位符 {product_name} / {customer} 等。"""
    if not isinstance(text, str):
        return text
    for k, v in context.items():
        text = text.replace("{" + k + "}", str(v))
    return text


def substitute_placeholders_deep(obj, context):
    """递归替换字典/列表中的占位符。"""
    if isinstance(obj, str):
        return substitute_placeholders(obj, context)
    if isinstance(obj, list):
        return [substitute_placeholders_deep(item, context) for item in obj]
    if isinstance(obj, dict):
        return {k: substitute_placeholders_deep(v, context) for k, v in obj.items()}
    return obj


# ============================================================
# FMEA 智能填充（auto_fill 模式）
# ============================================================

def auto_fill_failure_chain(chain: dict) -> dict:
    """自动填充失效链的 S/O/D 评分与 AP（基于模板 hint 字段）。

    当用户启用 auto_fill 模式时，从 chain 中的 _hint 字段读取建议值，
    并自动计算 AP（始终用 get_ap_priority 重算，确保与 S/O/D 一致，
    忽略模板预填的 ap_hint——因为模板 hint 可能存在不一致）。
    """
    s = chain.get("s_hint") if isinstance(chain.get("s_hint"), int) else chain.get("s", 5)
    o = chain.get("o_hint") if isinstance(chain.get("o_hint"), int) else chain.get("o", 5)
    d = chain.get("d_hint") if isinstance(chain.get("d_hint"), int) else chain.get("d", 5)
    # 始终用 get_ap_priority 重算，避免 ap_hint 与 s/o/d 组合不一致
    ap = get_ap_priority(s, o, d)

    chain["s"] = s
    chain["o"] = o
    chain["d"] = d
    chain["ap"] = ap
    return chain


def _guess_fill_value(left_val, first_col_val, col_idx, today_str, date_plus,
                      fmea_type, fmea_no, chain_index=None):
    """根据上下文猜测 ____ 应该填什么值。

    FMEA 智能填充规则（区别于 8D 的 5Why/6M 填充）：
    1. FMEA 团队表（第1列=角色名）：姓名/部门/分机号
    2. 优化措施表（第1列=序号）：责任人/截止日期/状态/措施后 SOD
    3. 表头信息表（左侧=字段名）：项目编号/批次号/编制人
    4. 措施后 S/O/D/AP 列：根据原 SOD 推断改进后值

    返回 None 表示无法判断，保留 ____
    """
    left_val_clean = (left_val or "").strip()
    first_col_clean = (first_col_val or "").strip()

    # 1. FMEA 团队表 / 签名栏：第1列是角色名
    if first_col_clean in ROLE_NAMES:
        name = ROLE_NAMES[first_col_clean]
        if name == "____":
            return None

        # 判断是团队表还是签名栏
        is_signature = any(first_col_clean.startswith(kw) for kw in ["编制", "审核", "批准"])

        if col_idx == 2:  # 姓名列
            return name
        if is_signature:
            # 签名栏：第3列=签名，第4列=日期
            if col_idx == 3:
                return name  # 签名 = 姓名
            if col_idx == 4:
                return today_str
        else:
            # 团队表：第3列=部门，第4列=联系方式
            if col_idx == 3:
                # 反查部门
                for dept, person in DEPT_PERSON.items():
                    if person == name:
                        return dept
                return "—"
            if col_idx == 4:
                return EXT_MAP.get(name, "____")

    # 2. 优化措施表：第1列是序号（1/2/3...）
    if first_col_clean.isdigit():
        seq = int(first_col_clean)
        # 序号行：按序号轮换责任人
        if col_idx == 5:  # 责任人列（参考 Sheet 6 列序：序号/FM/类型/描述/责任人/截止/状态/后S/后O/后D/后AP）
            return PERSON_ROTATION[(seq - 1) % len(PERSON_ROTATION)]
        if col_idx == 6:  # 截止日期列
            # AP=H 措施 7 天内、M 措施 14 天内、L 措施 30 天内
            # 默认按序号递增：7/10/14/21/30/45/60 天
            days_map = [7, 10, 14, 21, 30, 45, 60, 60, 60, 60]
            return date_plus(days_map[(seq - 1) % len(days_map)])
        if col_idx == 7:  # 状态列
            return "已建议"
        # 措施后 S/O/D/AP 列（8-11）：留 — 让用户实施后填写
        if col_idx in (8, 9, 10):
            return "—"
        if col_idx == 11:
            return "—"

    # 3. 表头信息表的空白字段
    if left_val_clean == "公司名称":
        return "（请填写公司名）"
    if left_val_clean == "项目编号":
        return f"P{today_str.replace('-', '')}-{1000 + (hash(first_col_val) % 9000)}"
    if left_val_clean == "FMEA 编号":
        return f"FMEA-{today_str.replace('-', '')}-001"
    if left_val_clean == "FMEA 开始日期":
        return today_str
    if left_val_clean == "FMEA 修订日期":
        return today_str
    if left_val_clean == "FMEA 编制人":
        return "张伟"
    if left_val_clean == "FMEA 审核人":
        return "王芳"
    if left_val_clean == "FMEA 批准人":
        return "孙健"
    if left_val_clean == "制造地址":
        return "（请填写制造地址）"
    if left_val_clean == "团队成员":
        return "张伟（推进者）、李娜（设计）、刘强（工艺）、陈静（质量）"

    # 4. 结构分析表的空白（PFMEA 工序要素）
    if left_val_clean in ["设备（M）", "机（Machine）"]:
        return "（请填写设备型号）"
    if left_val_clean in ["人员（M）", "人（Man）"]:
        return "（请填写岗位/资质要求）"
    if left_val_clean in ["材料（M）", "料（Material）"]:
        return "（请填写材料牌号）"
    if left_val_clean in ["方法（M）", "法（Method）"]:
        return "（请填写 SOP 编号）"
    if left_val_clean in ["环境（E）", "环（Environment）"]:
        return "（请填写温湿度要求）"
    if left_val_clean in ["测量（M）", "测（Measurement）"]:
        return "（请填写量具/MSA 编号）"

    # 5. 兜底：所有剩余 ____ 都填示例标记
    if left_val_clean and ("____" in left_val_clean or "请填写" in left_val_clean):
        return f"（示例数据，请替换为实际值）"

    return None


def _apply_auto_fill_xlsx(ws, fmea_type, fmea_no):
    """Excel 自动填充模式：扫描所有单元格，把 ____ 替换为合理示例值。

    FMEA 智能填充规则（区别于 8D 的 5Why/6M 填充）：
    - 团队表姓名：按角色分配化名
    - 优化措施责任人：按序号轮换分配
    - 日期字段：基于当前日期计算
    - 联系方式：内部分机号
    - S/O/D 评分：保留 ____ 让用户基于实际评估填写
    """
    today = datetime.date.today()
    today_str = today.strftime("%Y-%m-%d")
    date_plus = lambda days: (today + datetime.timedelta(days=days)).strftime("%Y-%m-%d")

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            if cell.value is None:
                continue
            val = str(cell.value).strip()
            if val != "____" and not val.startswith("____（") and not val.startswith("请填写"):
                continue

            left_cell = ws.cell(row=cell.row, column=cell.col_idx - 1) if cell.col_idx > 1 else None
            left_val = str(left_cell.value) if left_cell and left_cell.value else ""
            first_col_val = str(ws.cell(row=cell.row, column=1).value or "")

            replacement = _guess_fill_value(
                left_val=left_val,
                first_col_val=first_col_val,
                col_idx=cell.col_idx,
                today_str=today_str,
                date_plus=date_plus,
                fmea_type=fmea_type,
                fmea_no=fmea_no,
            )

            if replacement:
                cell.value = replacement


def _apply_auto_fill_word(doc, fmea_type, fmea_no):
    """Word 自动填充模式：扫描所有表格单元格，把 ____ 替换为合理示例值。"""
    today = datetime.date.today()
    today_str = today.strftime("%Y-%m-%d")
    date_plus = lambda days: (today + datetime.timedelta(days=days)).strftime("%Y-%m-%d")

    for table in doc.tables:
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells, start=1):
                text = cell.text.strip() if cell.text else ""
                if text != "____" and not text.startswith("____（") and not text.startswith("请填写"):
                    continue

                # 找同行左侧单元格
                left_val = ""
                if col_idx > 1 and col_idx <= len(row.cells):
                    left_val = (row.cells[col_idx - 2].text or "").strip()

                first_col_val = (row.cells[0].text or "").strip() if row.cells else ""

                replacement = _guess_fill_value(
                    left_val=left_val,
                    first_col_val=first_col_val,
                    col_idx=col_idx,
                    today_str=today_str,
                    date_plus=date_plus,
                    fmea_type=fmea_type,
                    fmea_no=fmea_no,
                )

                if replacement:
                    # 清空原段落，写入新值
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.text = ""
                    if cell.paragraphs:
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].text = replacement
                        else:
                            cell.paragraphs[0].add_run(replacement)


# ============================================================
# Excel 样式工具函数（FMEA 专用，区别于 8D 的样式）
# ============================================================

def get_thin_border():
    """获取细边框。"""
    side = Side(border_style="thin", color="808080")
    return Border(left=side, right=side, top=side, bottom=side)


def get_header_font():
    """FMEA 表头字体（白字深蓝底）。"""
    return Font(name=FONT_NAME, size=11, bold=True, color="FFFFFF")


def get_body_font():
    """FMEA 正文字体。"""
    return Font(name=FONT_NAME, size=10, color="000000")


def get_subheader_font():
    """FMEA 次级表头字体（白字亮蓝底）。"""
    return Font(name=FONT_NAME, size=10, bold=True, color="FFFFFF")


def get_header_fill():
    """FMEA 表头填充（深蓝）。"""
    return PatternFill(start_color=HEADER_FILL, end_color=HEADER_FILL, fill_type="solid")


def get_subheader_fill():
    """FMEA 次级表头填充（亮蓝）。"""
    return PatternFill(start_color=SUBHEADER_FILL, end_color=SUBHEADER_FILL, fill_type="solid")


def get_alt_fill():
    """FMEA 交替行填充（浅蓝）。"""
    return PatternFill(start_color=ALT_ROW_FILL, end_color=ALT_ROW_FILL, fill_type="solid")


def get_section_title_fill():
    """FMEA 章节标题填充（淡黄底）。"""
    return PatternFill(start_color=SECTION_TITLE_FILL, end_color=SECTION_TITLE_FILL, fill_type="solid")


def get_section_title_font():
    """FMEA 章节标题字体（深蓝字）。"""
    return Font(name=FONT_NAME, size=12, bold=True, color=HEADER_FILL)


def get_ap_fill(ap: str):
    """根据 AP 优先级返回对应的填充色。"""
    color_map = {
        "H": AP_H_FILL,
        "M": AP_M_FILL,
        "L": AP_L_FILL,
    }
    color = color_map.get(ap)
    if not color:
        return None
    return PatternFill(start_color=color, end_color=color, fill_type="solid")


def get_special_char_fill(sc: str):
    """根据特殊特性类型返回对应的填充色。"""
    if sc == "CC":
        return PatternFill(start_color=CC_FILL, end_color=CC_FILL, fill_type="solid")
    if sc == "SC":
        return PatternFill(start_color=SC_FILL, end_color=SC_FILL, fill_type="solid")
    return None


def get_status_fill(status: str):
    """根据措施状态返回对应的填充色。"""
    if not status:
        return None
    if status in ("已实施", "已关闭", "已完成"):
        return PatternFill(start_color=STATUS_DONE_FILL, end_color=STATUS_DONE_FILL, fill_type="solid")
    if status in ("已建议", "已决策", "进行中"):
        return PatternFill(start_color=STATUS_OPEN_FILL, end_color=STATUS_OPEN_FILL, fill_type="solid")
    return None


def apply_header_style(cell):
    """应用表头样式（深蓝底白字加粗居中）。"""
    cell.font = get_header_font()
    cell.fill = get_header_fill()
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_subheader_style(cell):
    """应用次级表头样式（亮蓝底白字加粗居中）。"""
    cell.font = get_subheader_font()
    cell.fill = get_subheader_fill()
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_body_style(cell, row_idx_in_data, is_alt_row=False):
    """应用正文样式（交替行浅蓝底）。"""
    cell.font = get_body_font()
    if is_alt_row:
        cell.fill = get_alt_fill()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_section_title_style(cell):
    """应用章节标题样式（淡黄底深蓝字）。"""
    cell.font = get_section_title_font()
    cell.fill = get_section_title_fill()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_ap_style(cell, ap: str):
    """应用 AP 单元格样式（按 H/M/L 高亮 + 加粗）。"""
    cell.font = Font(name=FONT_NAME, size=11, bold=True, color="000000")
    fill = get_ap_fill(ap)
    if fill:
        cell.fill = fill
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_special_char_style(cell, sc: str):
    """应用特殊特性单元格样式（CC 红底/SC 黄底加粗）。"""
    if sc:
        cell.font = Font(name=FONT_NAME, size=10, bold=True, color="C00000")
        fill = get_special_char_fill(sc)
        if fill:
            cell.fill = fill
    else:
        cell.font = get_body_font()
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_status_style(cell, status: str):
    """应用措施状态单元格样式。"""
    cell.font = get_body_font()
    fill = get_status_fill(status)
    if fill:
        cell.fill = fill
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


# ============================================================
# 行高自动计算（支持中文换行 + 合并单元格）
# ============================================================

def _display_width(s: str) -> int:
    """计算字符串的显示宽度（中文≈2，英文≈1），用于行高估算。"""
    w = 0
    for ch in str(s):
        if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f' or '\uff00' <= ch <= '\uffef':
            w += 2
        else:
            w += 1
    return w


def _calc_row_height(row_data, col_widths, line_height=17, min_height=20, max_height=409):
    """根据单元格内容和列宽精确计算行高，确保所有文字完全显示。

    核心逻辑：
    1. 遍历当前行每个单元格，计算该单元格文本在该列宽度下需要换几行
    2. 取所有单元格中最大的行数（即"最占行数的那一列"决定整行高度）
    3. 行高 = 最大行数 × 每行高度 + 上下边距

    CJK 字符宽度按英文 2 倍计算，显式换行符 \\n 也增加行数。

    Args:
        row_data: 该行各列的值列表
        col_widths: 各列宽度列表（Excel 列宽单位，约等于英文字符数）
        line_height: 每行高度（磅），默认 17
        min_height: 最小行高
        max_height: 最大行高（Excel 最大 409 磅）

    Returns:
        int: 计算出的行高（磅）
    """
    max_lines = 1
    for col_idx, value in enumerate(row_data):
        text = str(value) if value is not None else ""
        if not text:
            continue
        col_w = col_widths[col_idx] if col_idx < len(col_widths) else 10
        effective_width = max(col_w - 3, 4)

        segments = text.split('\n')
        total_lines = 0
        for segment in segments:
            if not segment.strip():
                total_lines += 1
                continue
            seg_display_w = _display_width(segment)
            lines_needed = -(-seg_display_w // effective_width)  # 向上取整
            lines_needed = max(1, lines_needed)
            total_lines += lines_needed

        max_lines = max(max_lines, total_lines)

    calculated_height = max_lines * line_height + 4
    return max(min_height, min(max_height, calculated_height))


def _recalc_all_row_heights(ws, line_height=17, min_height=20, max_height=409):
    """重新计算整个工作表所有行的行高（基于实际单元格内容和列宽）。

    这是在所有内容写入完成、auto_fill 执行完毕后调用的最终修正步骤。
    确保每个单元格的文字都完整显示，行高由该行文字最多的列决定。

    处理逻辑：
    1. 遍历每一行
    2. 对每个有内容的单元格，根据文本长度和列宽计算所需行数
    3. 合并单元格的宽度 = 涉及所有列宽之和
    4. 取该行所有单元格中最大行数，计算行高
    """
    for row_idx in range(1, ws.max_row + 1):
        max_lines = 1
        has_content = False
        for col_idx in range(1, ws.max_column + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            if cell.value is None:
                continue
            text = str(cell.value)
            if not text.strip():
                continue
            has_content = True

            col_letter = get_column_letter(col_idx)
            col_w = ws.column_dimensions[col_letter].width or 10

            # 合并单元格宽度累加
            effective_width = col_w
            for merge_range in ws.merged_cells.ranges:
                if cell.coordinate in merge_range:
                    total_w = 0
                    for c in range(merge_range.min_col, merge_range.max_col + 1):
                        cl = get_column_letter(c)
                        total_w += ws.column_dimensions[cl].width or 10
                    effective_width = total_w
                    break

            effective_width = max(effective_width - 3, 4)
            segments = text.split('\n')
            total_lines = 0
            for segment in segments:
                if not segment.strip():
                    total_lines += 1
                    continue
                seg_w = _display_width(segment)
                lines_needed = max(1, -(-seg_w // effective_width))
                total_lines += lines_needed
            max_lines = max(max_lines, max(1, total_lines))

        if has_content:
            calculated_height = max_lines * line_height + 4
            ws.row_dimensions[row_idx].height = max(min_height, min(max_height, calculated_height))


def set_column_widths(ws, widths):
    """设置列宽。"""
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


def write_section_title(ws, row, title, span_cols=6):
    """写入章节标题（淡黄底深蓝字）。"""
    ws.cell(row=row, column=1, value=title)
    if span_cols > 1:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span_cols)
    for col in range(1, span_cols + 1):
        apply_section_title_style(ws.cell(row=row, column=col))
    ws.row_dimensions[row].height = 28


def write_table(ws, start_row, headers, rows, col_widths,
                ap_col=None, special_char_col=None, status_col=None):
    """通用表格写入工具（支持 AP/特殊特性/状态列自动高亮）。

    Args:
        ws: 工作表
        start_row: 起始行
        headers: 表头列表
        rows: 数据行列表（每行是单元格值列表）
        col_widths: 列宽列表
        ap_col: AP 列的 1-indexed 位置（用于 H/M/L 高亮）
        special_char_col: 特殊特性列的 1-indexed 位置
        status_col: 状态列的 1-indexed 位置

    Returns:
        下一可用行号
    """
    # 设置列宽
    set_column_widths(ws, col_widths)

    # 写表头
    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=start_row, column=col_idx, value=h)
        apply_header_style(cell)
    ws.row_dimensions[start_row].height = 32

    # 写数据行
    for row_offset, row_data in enumerate(rows, start=1):
        current_row = start_row + row_offset
        is_alt = (row_offset % 2 == 1)

        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=value)

            # 优先级：AP 列 > 特殊特性列 > 状态列 > 普通正文
            if ap_col and col_idx == ap_col:
                apply_ap_style(cell, str(value) if value else "")
            elif special_char_col and col_idx == special_char_col:
                apply_special_char_style(cell, str(value) if value else "")
            elif status_col and col_idx == status_col:
                apply_status_style(cell, str(value) if value else "")
            else:
                apply_body_style(cell, row_offset, is_alt)

    return start_row + len(rows) + 1


def write_kv_block(ws, start_row, kv_pairs, label_col_width=22, value_col_width=50):
    """写入键值对块（左侧标签 + 右侧值）。"""
    set_column_widths(ws, [label_col_width, value_col_width])

    for offset, (k, v) in enumerate(kv_pairs):
        current_row = start_row + offset
        label_cell = ws.cell(row=current_row, column=1, value=k)
        value_cell = ws.cell(row=current_row, column=2, value=v)

        # 标签：次级表头样式
        apply_subheader_style(label_cell)
        # 值：正文样式
        apply_body_style(value_cell, offset + 1, (offset % 2 == 1))

    return start_row + len(kv_pairs) + 1


# ============================================================
# Excel 生成
# ============================================================

def create_excel_report(data: dict, output_path: Path):
    """生成 FMEA Excel 报告（7 Sheet）。"""
    wb = openpyxl.Workbook()

    chains = data.get("failure_chains", [])

    # ---------- Sheet 1: 表头 ----------
    ws1 = wb.active
    ws1.title = "1.表头"

    write_section_title(ws1, 1, "FMEA 表头信息", span_cols=2)

    header_info = [
        ("FMEA 类型", data.get("fmea_type", "")),
        ("公司名称", data.get("company", "____")),
        ("顾客名称", data.get("customer", "")),
        ("项目名称", data.get("product_name", "")),
        ("项目编号", data.get("project_no", "____")),
        ("系统层级", data.get("system_level", "—")),
        ("设计责任", data.get("design_responsibility", "—")),
        ("工艺名称", data.get("process_name", "—")),
        ("制造地址", data.get("manufacturing_site", "____")),
        ("团队成员", data.get("team_members", "____")),
        ("FMEA 开始日期", data.get("start_date", datetime.date.today().isoformat())),
        ("FMEA 修订日期", datetime.date.today().isoformat()),
        ("FMEA 编号", data.get("fmea_no", f"FMEA-{datetime.date.today().strftime('%Y%m%d')}-001")),
        ("FMEA 编制人", "____"),
        ("FMEA 审核人", "____"),
        ("FMEA 批准人", "____"),
        ("使用模板", data.get("template", "generic-fmea")),
        ("参考标准", "AIAG & VDA FMEA Handbook 2019"),
    ]
    write_kv_block(ws1, 2, header_info, label_col_width=22, value_col_width=50)

    # ---------- Sheet 2: 结构分析 ----------
    ws2 = wb.create_sheet("2.结构分析")
    write_section_title(ws2, 1, "步骤二：结构分析", span_cols=3)

    structure = data.get("structure_tree", [])
    if structure:
        rows = [
            [item.get("level", ""), item.get("name", ""), item.get("description", "")]
            for item in structure
        ]
        write_table(ws2, 2, ["层级", "要素名称", "说明"], rows, [15, 35, 50])
    else:
        # 占位提示
        ws2.cell(row=2, column=1, value="（请在 FMEA 团队会议中填写结构树/过程流程图）")
        ws2.merge_cells("A2:C2")
        apply_body_style(ws2.cell(row=2, column=1), 1, False)

    # ---------- Sheet 3: 功能分析 ----------
    ws3 = wb.create_sheet("3.功能分析")
    write_section_title(ws3, 1, "步骤三：功能分析", span_cols=3)

    functions = data.get("function_tree", [])
    if functions:
        rows = [
            [item.get("level", ""), item.get("name", ""), item.get("function", "")]
            for item in functions
        ]
        write_table(ws3, 2, ["层级", "要素名称", "功能描述"], rows, [15, 25, 55])
    else:
        ws3.cell(row=2, column=1, value="（功能定义遵循「主动动词 + 可测量名词」格式）")
        ws3.merge_cells("A2:C2")
        apply_body_style(ws3.cell(row=2, column=1), 1, False)

    # ---------- Sheet 4: 失效分析 ----------
    ws4 = wb.create_sheet("4.失效分析")
    write_section_title(ws4, 1, "步骤四：失效分析（失效链 FE→FM→FC）", span_cols=4)

    chain_rows = [
        [i, c.get("fe", ""), c.get("fm", ""), c.get("fc", "")]
        for i, c in enumerate(chains, start=1)
    ]
    write_table(ws4, 2, ["序号", "失效影响（FE）", "失效模式（FM）", "失效起因（FC）"],
                chain_rows, [6, 45, 35, 45])

    # ---------- Sheet 5: 风险分析 ----------
    ws5 = wb.create_sheet("5.风险分析")
    write_section_title(ws5, 1, "步骤五：风险分析（S/O/D 评级 + AP + 特殊特性）", span_cols=10)

    risk_rows = []
    for i, c in enumerate(chains, start=1):
        s = c.get("s", 5)
        o = c.get("o", 5)
        d = c.get("d", 5)
        ap = c.get("ap", get_ap_priority(s, o, d))
        sc = get_special_characteristic(s, ap)
        risk_rows.append([
            i,
            c.get("fe", ""),
            c.get("fm", ""),
            c.get("pc", ""),
            c.get("dc", ""),
            s, o, d, ap, sc,
        ])

    write_table(ws5, 2,
                ["序号", "失效影响（FE）", "失效模式（FM）", "预防控制（PC）", "探测控制（DC）",
                 "S", "O", "D", "AP", "特殊特性"],
                risk_rows,
                [6, 30, 25, 35, 35, 6, 6, 6, 6, 10],
                ap_col=9,           # AP 列高亮
                special_char_col=10)  # 特殊特性列高亮

    # ---------- Sheet 6: 优化措施 ----------
    ws6 = wb.create_sheet("6.优化措施")
    write_section_title(ws6, 1, "步骤六：优化措施（PC/DC 改进 + 措施跟踪）", span_cols=11)

    measures = data.get("optimization_measures", [])
    measure_rows = []
    for i, m in enumerate(measures, start=1):
        measure_rows.append([
            i,
            m.get("fm", ""),
            m.get("type", "PC+DC 改进"),
            m.get("description", ""),
            m.get("owner", "____"),
            m.get("due_date", "____"),
            m.get("status", "已建议"),
            "—",  # 措施后 S
            "—",  # 措施后 O
            "—",  # 措施后 D
            "—",  # 措施后 AP
        ])

    write_table(ws6, 2,
                ["序号", "失效模式", "措施类型", "措施描述", "责任人", "截止日期", "状态",
                 "措施后 S", "措施后 O", "措施后 D", "措施后 AP"],
                measure_rows,
                [6, 30, 12, 50, 12, 14, 12, 9, 9, 9, 10],
                status_col=7)

    # ---------- Sheet 7: 风险矩阵 ----------
    ws7 = wb.create_sheet("7.风险矩阵")
    write_section_title(ws7, 1, "步骤七：风险矩阵（S×O 热力图 + 风险统计）", span_cols=12)

    # S×O 矩阵（10x10）
    ws7.cell(row=3, column=1, value="S\\O")
    apply_header_style(ws7.cell(row=3, column=1))

    for o in range(1, 11):
        c = ws7.cell(row=3, column=o + 1, value=o)
        apply_header_style(c)

    for s in range(10, 0, -1):
        row_idx = 14 - s
        c = ws7.cell(row=row_idx, column=1, value=s)
        apply_header_style(c)
        for o in range(1, 11):
            ap = get_ap_priority(s, o, 5)  # D=5 中位数
            cell = ws7.cell(row=row_idx, column=o + 1, value=ap)
            apply_ap_style(cell, ap)

    # 图例
    ws7.cell(row=16, column=1, value="图例：").font = Font(name=FONT_NAME, bold=True)
    for col, (label, ap) in enumerate([(2, "H"), (3, "M"), (4, "L")], start=2):
        c = ws7.cell(row=16, column=col, value=f"{label}（{'高' if ap=='H' else '中' if ap=='M' else '低'}）")
        apply_ap_style(c, ap)

    # 风险统计
    ws7.cell(row=18, column=1, value="风险统计").font = Font(name=FONT_NAME, bold=True, size=12)
    h_count = sum(1 for c in chains if c.get("ap") == "H")
    m_count = sum(1 for c in chains if c.get("ap") == "M")
    l_count = sum(1 for c in chains if c.get("ap") == "L")
    cc_count = sum(1 for c in chains if c.get("s", 0) >= 9)
    sc_count = sum(1 for c in chains if c.get("s", 0) == 8 and c.get("ap") in ("H", "M"))

    stats = [
        ("失效链总数", len(chains)),
        ("AP=H（高优先级）", h_count),
        ("AP=M（中优先级）", m_count),
        ("AP=L（低优先级）", l_count),
        ("CC（关键特性，S≥9）", cc_count),
        ("SC（特殊特性，S=8 且 AP=H/M）", sc_count),
    ]
    stats_start = 19
    for i, (k, v) in enumerate(stats):
        r = stats_start + i
        ws7.cell(row=r, column=1, value=k).font = Font(name=FONT_NAME, bold=True)
        cell = ws7.cell(row=r, column=2, value=v)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.font = Font(name=FONT_NAME, bold=True, size=12, color=HEADER_FILL)

    # 设置各 Sheet 列宽
    set_column_widths(ws7, [22] + [6] * 11)

    # ---------- 自动填充 ----------
    if data.get("auto_fill"):
        for ws in [ws1, ws2, ws3, ws4, ws5, ws6, ws7]:
            _apply_auto_fill_xlsx(ws, data.get("fmea_type", ""), data.get("fmea_no", ""))

    # ---------- 行高自动计算 ----------
    for ws in [ws1, ws2, ws3, ws4, ws5, ws6, ws7]:
        _recalc_all_row_heights(ws)

    # 保存
    wb.save(output_path)
    print(f"[OK] Excel 报告已生成：{output_path}")


# ============================================================
# Word 生成
# ============================================================

# Word 字体辅助函数
def set_cell_bg(cell, color_hex):
    """设置 Word 表格单元格背景色。"""
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    """设置 Word 表格单元格四边边框。"""
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '808080')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_run_font(run, font_name=FONT_NAME, size=10.5, bold=False, color=None):
    """设置 Word run 的字体属性。"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 设置中文字体
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        from docx.oxml import OxmlElement
        r_fonts = OxmlElement('w:rFonts')
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn('w:eastAsia'), font_name)


def add_paragraph(doc, text, bold=False, size=10.5, color=None, alignment=None, indent=False):
    """添加段落。"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    """添加标题。"""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    size_map = {0: 22, 1: 16, 2: 13, 3: 11}
    set_run_font(run, size=size_map.get(level, 11), bold=True, color=HEADER_FILL)
    return h


def add_table(doc, headers, rows, col_widths_cm=None,
              ap_col=None, special_char_col=None, status_col=None):
    """添加表格（支持 AP/特殊特性/状态列高亮）。

    Args:
        doc: Document 对象
        headers: 表头列表
        rows: 数据行列表
        col_widths_cm: 列宽（厘米）列表
        ap_col: AP 列的 1-indexed 位置
        special_char_col: 特殊特性列的 1-indexed 位置
        status_col: 状态列的 1-indexed 位置
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 设置列宽
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # 写表头
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        # 清空原段落
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_font(run, size=10.5, bold=True, color="FFFFFF")
        set_cell_bg(hdr[i], HEADER_FILL)
        set_cell_borders(hdr[i])
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 写数据行
    for row_idx, row_data in enumerate(rows, start=1):
        is_alt = (row_idx % 2 == 1)
        cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row_data, start=1):
            value_str = str(value) if value is not None else ""
            cell = cells[col_idx - 1]
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ""
            run = cell.paragraphs[0].add_run(value_str)

            # 按列类型应用样式
            if ap_col and col_idx == ap_col:
                set_run_font(run, size=11, bold=True, color="000000")
                ap_val = value_str
                if ap_val == "H":
                    set_cell_bg(cell, AP_H_FILL)
                elif ap_val == "M":
                    set_cell_bg(cell, AP_M_FILL)
                elif ap_val == "L":
                    set_cell_bg(cell, AP_L_FILL)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif special_char_col and col_idx == special_char_col:
                if value_str:
                    set_run_font(run, size=10.5, bold=True, color="C00000")
                    if value_str == "CC":
                        set_cell_bg(cell, CC_FILL)
                    elif value_str == "SC":
                        set_cell_bg(cell, SC_FILL)
                else:
                    set_run_font(run, size=10.5)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif status_col and col_idx == status_col:
                set_run_font(run, size=10.5)
                if value_str in ("已实施", "已关闭", "已完成"):
                    set_cell_bg(cell, STATUS_DONE_FILL)
                elif value_str in ("已建议", "已决策", "进行中"):
                    set_cell_bg(cell, STATUS_OPEN_FILL)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_run_font(run, size=10.5)
                if is_alt:
                    set_cell_bg(cell, ALT_ROW_FILL)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            set_cell_borders(cell)

    return table


def add_kv_table(doc, kv_pairs, label_width_cm=4.5, value_width_cm=12):
    """添加键值对表格（左侧标签 + 右侧值）。"""
    table = doc.add_table(rows=len(kv_pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (k, v) in enumerate(kv_pairs):
        cells = table.rows[i].cells
        cells[0].width = Cm(label_width_cm)
        cells[1].width = Cm(value_width_cm)

        # 标签
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.text = ""
        run = cells[0].paragraphs[0].add_run(k)
        set_run_font(run, size=10.5, bold=True, color="FFFFFF")
        set_cell_bg(cells[0], SUBHEADER_FILL)
        set_cell_borders(cells[0])

        # 值
        for p in cells[1].paragraphs:
            for r in p.runs:
                r.text = ""
        run = cells[1].paragraphs[0].add_run(str(v))
        set_run_font(run, size=10.5)
        if i % 2 == 1:
            set_cell_bg(cells[1], ALT_ROW_FILL)
        set_cell_borders(cells[1])

    return table


def set_doc_default_font(doc):
    """设置文档默认字体为微软雅黑。"""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def set_page_header(doc, fmea_no):
    """设置页眉显示 FMEA 编号。"""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"FMEA 编号：{fmea_no}")
    set_run_font(run, size=9, color="808080")


def create_docx_report(data: dict, output_path: Path):
    """生成 FMEA Word 报告（纯表格版本，与 xlsx 7 Sheet 内容完全一致）。

    用户要求：docx 和 xlsx 都应该是表格内容，内容应该是一样的。
    因此 docx 不再是"七步法报告"（段落+表格混合），而是 7 个表格（对应 xlsx 的 7 个 Sheet）：
    - 表 1：表头信息（项目信息表）
    - 表 2：结构分析表
    - 表 3：功能分析表
    - 表 4：失效分析表（FE→FM→FC）
    - 表 5：风险分析表（S/O/D + AP + CC/SC）
    - 表 6：优化措施表
    - 表 7：风险矩阵表 + 风险统计表
    """
    doc = Document()
    set_doc_default_font(doc)

    fmea_type = data.get("fmea_type", "FMEA")
    fmea_no = data.get("fmea_no", f"FMEA-{datetime.date.today().strftime('%Y%m%d')}-001")
    set_page_header(doc, fmea_no)

    # 文档标题
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{fmea_type} 分析表")
    set_run_font(run, size=22, bold=True, color=HEADER_FILL)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"产品：{data.get('product_name', '')}    客户：{data.get('customer', '')}    日期：{datetime.date.today().isoformat()}"
    )
    set_run_font(run, size=11)

    chains = data.get("failure_chains", [])

    # ===== 表 1：表头信息（对应 xlsx Sheet 1） =====
    add_heading(doc, "表 1  表头信息", level=1)
    header_info = [
        ("FMEA 类型", data.get("fmea_type", "")),
        ("公司名称", data.get("company", "")),
        ("顾客名称", data.get("customer", "")),
        ("项目名称", data.get("product_name", "")),
        ("项目编号", data.get("project_no", "")),
        ("系统层级", data.get("system_level", "—")),
        ("设计责任", data.get("design_responsibility", "—")),
        ("工艺名称", data.get("process_name", "—")),
        ("制造地址", data.get("manufacturing_site", "")),
        ("团队成员", data.get("team_members", "")),
        ("FMEA 开始日期", data.get("start_date", datetime.date.today().isoformat())),
        ("FMEA 修订日期", datetime.date.today().isoformat()),
        ("FMEA 编号", fmea_no),
        ("FMEA 编制人", data.get("preparer", "")),
        ("FMEA 审核人", data.get("reviewer", "")),
        ("FMEA 批准人", data.get("approver", "")),
        ("使用模板", data.get("template", "generic-fmea")),
        ("参考标准", "AIAG & VDA FMEA Handbook 2019"),
    ]
    add_kv_table(doc, header_info, label_width_cm=4.5, value_width_cm=12)

    # ===== 表 2：结构分析（对应 xlsx Sheet 2） =====
    add_heading(doc, "表 2  结构分析（步骤二）", level=1)
    structure = data.get("structure_tree", [])
    if structure:
        struct_rows = [
            [item.get("level", ""), item.get("name", ""), item.get("description", "")]
            for item in structure
        ]
        add_table(doc, ["层级", "要素名称", "说明"], struct_rows,
                  col_widths_cm=[2.5, 4, 9])
    else:
        add_paragraph(doc, "（请在 FMEA 团队会议中填写结构树/过程流程图。DFMEA 用结构树或方块图，PFMEA 用过程流程图 + 4M1E 工作要素。）", indent=True)

    # ===== 表 3：功能分析（对应 xlsx Sheet 3） =====
    add_heading(doc, "表 3  功能分析（步骤三）", level=1)
    functions = data.get("function_tree", [])
    if functions:
        func_rows = [
            [item.get("level", ""), item.get("name", ""), item.get("function", "")]
            for item in functions
        ]
        add_table(doc, ["层级", "要素名称", "功能描述"], func_rows,
                  col_widths_cm=[2.5, 4, 9])
    else:
        add_paragraph(doc, "（功能定义遵循「主动动词 + 可测量名词」格式。DFMEA 用功能树或参数图 P-图，PFMEA 用过程功能分析。）", indent=True)

    # ===== 表 4：失效分析（对应 xlsx Sheet 4） =====
    add_heading(doc, "表 4  失效分析（步骤四：失效链 FE→FM→FC）", level=1)
    if chains:
        chain_rows = [
            [i, c.get("fe", ""), c.get("fm", ""), c.get("fc", "")]
            for i, c in enumerate(chains, start=1)
        ]
        add_table(doc, ["序号", "失效影响（FE）", "失效模式（FM）", "失效起因（FC）"],
                  chain_rows, col_widths_cm=[1.2, 5, 4, 5])
    else:
        add_paragraph(doc, "（无失效链数据）", indent=True)

    # ===== 表 5：风险分析（对应 xlsx Sheet 5） =====
    add_heading(doc, "表 5  风险分析（步骤五：S/O/D 评级 + AP + 特殊特性）", level=1)
    if chains:
        risk_rows = []
        for i, c in enumerate(chains, start=1):
            s = c.get("s", 5)
            o = c.get("o", 5)
            d = c.get("d", 5)
            ap = c.get("ap", get_ap_priority(s, o, d))
            sc = get_special_characteristic(s, ap)
            risk_rows.append([
                i,
                c.get("fe", ""),
                c.get("fm", ""),
                c.get("pc", ""),
                c.get("dc", ""),
                s, o, d, ap, sc,
            ])
        add_table(doc,
                  ["序号", "失效影响（FE）", "失效模式（FM）", "预防控制（PC）", "探测控制（DC）",
                   "S", "O", "D", "AP", "特殊特性"],
                  risk_rows,
                  col_widths_cm=[1, 3.5, 2.5, 2.5, 2.5, 0.8, 0.8, 0.8, 1, 1.5],
                  ap_col=9, special_char_col=10)

    # ===== 表 6：优化措施（对应 xlsx Sheet 6） =====
    add_heading(doc, "表 6  优化措施（步骤六：PC/DC 改进 + 措施跟踪）", level=1)
    measures = data.get("optimization_measures", [])
    if measures:
        measure_rows = []
        for i, m in enumerate(measures, start=1):
            measure_rows.append([
                i,
                m.get("fm", ""),
                m.get("type", "PC+DC 改进"),
                m.get("description", ""),
                m.get("owner", ""),
                m.get("due_date", ""),
                m.get("status", "已建议"),
                "—",  # 措施后 S
                "—",  # 措施后 O
                "—",  # 措施后 D
                "—",  # 措施后 AP
            ])
        add_table(doc,
                  ["序号", "失效模式", "措施类型", "措施描述", "责任人", "截止日期", "状态",
                   "措施后 S", "措施后 O", "措施后 D", "措施后 AP"],
                  measure_rows,
                  col_widths_cm=[1, 3, 1.5, 4, 1.2, 1.5, 1.2, 1, 1, 1, 1.2],
                  status_col=7)
    else:
        add_paragraph(doc, "（无优化措施数据，所有失效链 AP 均为 L）", indent=True)

    # ===== 表 7：风险矩阵 + 风险统计（对应 xlsx Sheet 7） =====
    add_heading(doc, "表 7  风险矩阵（S×O 热力图，D=5 中位数）", level=1)

    # 7.1 S×O 矩阵表（10x10）
    matrix_headers = ["S\\O"] + [str(o) for o in range(1, 11)]
    matrix_rows = []
    for s in range(10, 0, -1):
        row = [str(s)]
        for o in range(1, 11):
            ap = get_ap_priority(s, o, 5)
            row.append(ap)
        matrix_rows.append(row)
    add_table(doc, matrix_headers, matrix_rows,
              col_widths_cm=[1] + [0.9]*10,
              ap_col=None)  # AP 列已经是字符串，不需要再高亮（用通用样式）

    # 7.2 风险统计表
    add_heading(doc, "表 7-2  风险统计", level=2)
    h_count = sum(1 for c in chains if c.get("ap") == "H")
    m_count = sum(1 for c in chains if c.get("ap") == "M")
    l_count = sum(1 for c in chains if c.get("ap") == "L")
    cc_count = sum(1 for c in chains if c.get("s", 0) >= 9)
    sc_count = sum(1 for c in chains if c.get("s", 0) == 8 and c.get("ap") in ("H", "M"))

    stats_rows = [
        ["失效链总数", str(len(chains))],
        ["AP=H（高优先级）", str(h_count)],
        ["AP=M（中优先级）", str(m_count)],
        ["AP=L（低优先级）", str(l_count)],
        ["CC（关键特性，S≥9）", str(cc_count)],
        ["SC（特殊特性，S=8 且 AP=H/M）", str(sc_count)],
    ]
    add_table(doc, ["统计项", "数量"], stats_rows,
              col_widths_cm=[8, 3])

    # 7.3 图例
    add_paragraph(doc, "")
    add_paragraph(doc, "图例：H=高优先级（红）  M=中优先级（黄）  L=低优先级（绿）", indent=True)

    # ===== 签名栏 =====
    add_heading(doc, "签名栏", level=1)
    sign_rows = [
        ["编制（FMEA 推进者）", "____", "____", "____"],
        ["审核（质量经理）", "____", "____", "____"],
        ["批准（项目经理）", "____", "____", "____"],
    ]
    add_table(doc, ["角色", "姓名", "签名", "日期"], sign_rows,
              col_widths_cm=[5, 3, 3, 3])

    # ---------- 自动填充 ----------
    if data.get("auto_fill"):
        _apply_auto_fill_word(doc, fmea_type, fmea_no)

    # 保存
    doc.save(output_path)
    print(f"[OK] Word 报告已生成：{output_path}")


# ============================================================
# 主入口
# ============================================================

def parse_args():
    parser = argparse.ArgumentParser(description="PFMEA/DFMEA 报告生成器")
    parser.add_argument("--fmea-type", required=True, choices=["DFMEA", "PFMEA"], help="FMEA 类型")
    parser.add_argument("--product", required=True, help="产品名称")
    parser.add_argument("--customer", required=True, help="客户名称")
    parser.add_argument("--project-no", default="", help="项目编号")
    parser.add_argument("--system-level", default="", help="系统层级（DFMEA）")
    parser.add_argument("--design-responsibility", default="", help="设计责任（DFMEA）")
    parser.add_argument("--process-name", default="", help="工艺名称（PFMEA）")
    parser.add_argument("--process-steps", default="", help="工序清单（逗号分隔，PFMEA）")
    parser.add_argument("--manufacturing-site", default="", help="制造地址（PFMEA）")
    parser.add_argument("--team", default="", help="团队成员")
    parser.add_argument("--template", default="generic-fmea", help="模板 slug")
    parser.add_argument("--auto-fill", action="store_true", help="自动填充示例值（S/O/D/责任人/日期等）")
    parser.add_argument("--output-dir", default=".", help="输出目录")
    parser.add_argument("--failure-chains-json", default="", help="自定义失效链 JSON 文件路径")
    return parser.parse_args()


def main():
    args = parse_args()

    # 模板目录（脚本所在目录的上级 templates/）
    script_dir = Path(__file__).resolve().parent
    templates_dir = script_dir.parent / "templates"

    # 加载模板
    template = load_template(args.template, templates_dir)
    print(f"[INFO] 使用模板：{template.get('name', args.template)}")

    # 构建上下文
    context = {
        "product_name": args.product,
        "customer": args.customer,
        "process_name": args.process_name,
    }

    # 失效链
    chains = template.get("failure_chains_template", [])
    if args.failure_chains_json:
        with open(args.failure_chains_json, "r", encoding="utf-8") as f:
            chains = json.load(f)

    # 占位符替换
    chains = substitute_placeholders_deep(chains, context)

    # 自动填充
    if args.auto_fill:
        chains = [auto_fill_failure_chain(c) for c in chains]
        print(f"[INFO] 已自动填充 {len(chains)} 条失效链的 S/O/D 评分")
    else:
        # 仍计算 AP（如果 S/O/D 已提供）
        for c in chains:
            s = c.get("s", c.get("s_hint", 5))
            o = c.get("o", c.get("o_hint", 5))
            d = c.get("d", c.get("d_hint", 5))
            if isinstance(s, int) and isinstance(o, int) and isinstance(d, int):
                c["s"] = s
                c["o"] = o
                c["d"] = d
                c["ap"] = c.get("ap_hint") or get_ap_priority(s, o, d)

    # SOD 一致性检查（仅警告，不阻止生成）
    warnings_count = 0
    for i, c in enumerate(chains, start=1):
        warnings = check_sod_consistency(c)
        if warnings:
            warnings_count += len(warnings)
            for w in warnings:
                print(f"[WARN] 失效链 #{i}: {w}")

    # 构建报告数据
    fmea_no = f"FMEA-{datetime.date.today().strftime('%Y%m%d')}-001"
    data = {
        "fmea_type": args.fmea_type,
        "company": "____",
        "customer": args.customer,
        "product_name": args.product,
        "project_no": args.project_no or "____",
        "system_level": args.system_level,
        "design_responsibility": args.design_responsibility,
        "process_name": args.process_name,
        "manufacturing_site": args.manufacturing_site or "____",
        "team_members": args.team or "____",
        "template": args.template,
        "fmea_no": fmea_no,
        "preparer": "____",
        "reviewer": "____",
        "approver": "____",
        "auto_fill": args.auto_fill,
        "structure_tree": [],  # 可扩展：从模板提取
        "function_tree": [],   # 可扩展：从模板提取
        "failure_chains": chains,
        "optimization_measures": [
            {
                "fm": c.get("fm", ""),
                "type": "PC+DC 改进",
                "description": c.get("pc", "") + " | " + c.get("dc", ""),
                "owner": "____",
                "due_date": "____",
                "status": "已建议",
            }
            for c in chains
            if c.get("ap") in ("H", "M")
        ],
    }

    # 输出路径
    output_dir = Path(args.output_dir).expanduser()
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    xlsx_path = output_dir / f"FMEA_{args.fmea_type}_{timestamp}.xlsx"
    docx_path = output_dir / f"FMEA_{args.fmea_type}_{timestamp}.docx"

    # 生成报告
    create_excel_report(data, xlsx_path)
    create_docx_report(data, docx_path)

    print(f"\n{'=' * 60}")
    print(f"✅ FMEA 报告生成完成！")
    print(f"{'=' * 60}")
    print(f"  Excel: {xlsx_path}")
    print(f"  Word:  {docx_path}")
    print(f"\n  失效链数: {len(chains)}")
    h_count = sum(1 for c in chains if c.get("ap") == "H")
    m_count = sum(1 for c in chains if c.get("ap") == "M")
    l_count = sum(1 for c in chains if c.get("ap") == "L")
    print(f"  AP=H: {h_count}  AP=M: {m_count}  AP=L: {l_count}")
    if warnings_count:
        print(f"  ⚠️  SOD 一致性警告: {warnings_count} 条（详见上方 [WARN]）")
    if args.auto_fill:
        print(f"  ✨ 已启用 auto_fill：表头/团队/优化措施表的 ____ 已自动填充示例值")


if __name__ == "__main__":
    main()
